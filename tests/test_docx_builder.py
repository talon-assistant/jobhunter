"""Tests for docx_builder.py."""

from pathlib import Path

from jobhunter.core.docx_builder import _split_bold_lead, build_cover_letter, build_resume


def test_split_bold_lead_colon():
    bold, rest = _split_bold_lead("Led team: managed 12 engineers across 3 sites")
    assert bold == "Led team:"
    assert rest == "managed 12 engineers across 3 sites"


def test_split_bold_lead_preposition():
    bold, rest = _split_bold_lead("Reduced security incidents by 45% through zero trust")
    assert "by" in bold or "through" in bold
    assert rest  # should have something


def test_split_bold_lead_short_text():
    bold, rest = _split_bold_lead("Short text")
    assert bold == "Short text"
    assert rest == ""


def test_build_resume(tmp_path):
    sections = {
        "Experience": [
            "Led team of 12 engineers: delivering security solutions across 3 SOCs",
            "Reduced incidents by 45% through zero trust implementation",
        ],
        "Skills": [
            "CISSP, CISM, AWS Security Specialty",
        ],
    }
    out = build_resume(
        sections,
        name="John Doe",
        email="john@example.com",
        output_path=tmp_path / "test_resume.docx",
    )
    assert out.exists()
    assert out.suffix == ".docx"

    # Verify it's a valid DOCX
    from docx import Document
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert any("JOHN DOE" in t for t in texts)
    assert any("EXPERIENCE" in t for t in texts)


def test_build_cover_letter(tmp_path):
    text = (
        "Dear Hiring Manager,\n\n"
        "I am writing about the position.\n\n"
        "My experience aligns well with your needs.\n\n"
        "Thank you for your consideration."
    )
    out = build_cover_letter(
        text,
        name="John Doe",
        output_path=tmp_path / "test_letter.docx",
    )
    assert out.exists()
    assert out.suffix == ".docx"


def test_build_resume_applies_template_styles(tmp_path):
    from docx import Document
    from jobhunter.templates import TEMPLATES

    sections = {"Experience": ["Led a 6-person team, cutting MTTR by 40%"]}

    for key in ("minimal", "technical", "executive"):
        out = build_resume(
            sections,
            name="Jane Doe",
            output_path=tmp_path / f"resume_{key}.docx",
            template=key,
        )
        assert out.exists()
        doc = Document(str(out))
        name_run = doc.paragraphs[0].runs[0]
        assert name_run.font.name == TEMPLATES[key]["style"]["font"]

    # Minimal ATS renders black headings and no rule paragraphs
    doc = Document(str(tmp_path / "resume_minimal.docx"))
    assert all(p.text or not p.runs for p in doc.paragraphs)


def test_build_resume_unknown_template_falls_back(tmp_path):
    out = build_resume(
        {"Skills": ["Python, Qt, SQLite"]},
        name="Jane Doe",
        output_path=tmp_path / "resume_fallback.docx",
        template="does-not-exist",
    )
    assert out.exists()
