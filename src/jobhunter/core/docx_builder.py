"""Build resume and cover letter DOCX files using python-docx."""

from __future__ import annotations

import logging
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

log = logging.getLogger(__name__)


def _add_horizontal_rule(doc, color: str = "CCCCCC"):
    """Add a thin horizontal rule paragraph under a heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): color,
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _split_bold_lead(text: str) -> tuple[str, str]:
    """Split a bullet into (bold_lead, rest) for formatting.

    Tries several split strategies in order of quality:
    1. Colon-space (strongest signal)
    2. Period-space (for sentence-lead bullets)
    3. Strong prepositions (by, through, across, resulting in)
    4. Comma
    5. Weak prepositions (for, with, to)
    6. Fallback: first 5 words
    """
    # Strategy 1: colon-space
    if ": " in text:
        idx = text.index(": ")
        return text[: idx + 1], text[idx + 2 :]

    # Strategy 2: period-space (but not after abbreviations like "Inc." or numbers)
    period_match = re.search(r"(?<![A-Z])\. ", text)
    if period_match and period_match.start() > 10:
        idx = period_match.start()
        return text[: idx + 1], text[idx + 2 :]

    # Strategy 3: strong prepositions
    for prep in (" by ", " through ", " across ", " resulting in "):
        if prep in text:
            idx = text.index(prep)
            if idx > 15:
                return text[: idx + len(prep)].rstrip(), text[idx + len(prep) :]

    # Strategy 4: comma
    if ", " in text:
        idx = text.index(", ")
        if idx > 15:
            return text[: idx + 1], text[idx + 2 :]

    # Strategy 5: weak prepositions
    for prep in (" for ", " with ", " to "):
        if prep in text:
            idx = text.index(prep)
            if idx > 15:
                return text[: idx + len(prep)].rstrip(), text[idx + len(prep) :]

    # Fallback: first 5 words
    words = text.split()
    if len(words) > 5:
        lead = " ".join(words[:5])
        rest = " ".join(words[5:])
        return lead, rest

    return text, ""


def build_resume(
    sections: dict[str, list[str]],
    *,
    name: str = "",
    email: str = "",
    phone: str = "",
    location: str = "",
    output_path: str | Path,
    template: str = "classic",
) -> Path:
    """Build a resume DOCX from selected bullets organized by section.

    Parameters
    ----------
    sections : dict mapping section name -> list of bullet strings
    name, email, phone, location : header contact info
    output_path : where to save the DOCX
    template : key from jobhunter.templates.TEMPLATES (font, sizes,
        accent color, margins, alignment). Unknown keys fall back to
        "classic".

    Returns the output Path.
    """
    from jobhunter.templates import get_template_style

    style = get_template_style(template)
    font = style["font"]
    accent = RGBColor(*style["accent"])
    accent_hex = "%02X%02X%02X" % style["accent"]
    header_align = (
        WD_ALIGN_PARAGRAPH.CENTER if style["center_header"]
        else WD_ALIGN_PARAGRAPH.LEFT
    )

    doc = Document()

    # Page margins
    top, bottom, left, right = style["margins"]
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

    # Header: Name
    if name:
        heading = doc.add_paragraph()
        heading.alignment = header_align
        run = heading.add_run(name.upper())
        run.bold = True
        run.font.size = Pt(style["name_size"])
        run.font.name = font
        run.font.color.rgb = accent

    # Header: Contact line
    contact_parts = [p for p in [email, phone, location] if p]
    if contact_parts:
        contact = doc.add_paragraph()
        contact.alignment = header_align
        run = contact.add_run(" | ".join(contact_parts))
        run.font.size = Pt(10)
        run.font.name = font

    if style["rules"]:
        _add_horizontal_rule(doc, accent_hex)

    # Sections with bullets
    body_size = Pt(style["body_size"])
    for section_name, bullets in sections.items():
        if not bullets:
            continue

        # Section header
        header = doc.add_paragraph()
        run = header.add_run(section_name.upper())
        run.bold = True
        run.font.size = Pt(style["heading_size"])
        run.font.name = font
        run.font.color.rgb = accent
        header.paragraph_format.space_before = Pt(8)
        header.paragraph_format.space_after = Pt(2)

        if style["rules"]:
            _add_horizontal_rule(doc, "CCCCCC")

        # Bullets
        for bullet_text in bullets:
            p = doc.add_paragraph(style="List Bullet")
            bold_lead, rest = _split_bold_lead(bullet_text)

            if rest:
                run_bold = p.add_run(bold_lead + " ")
                run_bold.bold = True
                run_bold.font.size = body_size
                run_bold.font.name = font

                run_rest = p.add_run(rest)
                run_rest.font.size = body_size
                run_rest.font.name = font
            else:
                run = p.add_run(bold_lead)
                run.font.size = body_size
                run.font.name = font

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    log.info("Resume saved to %s (template=%s)", out, template)
    return out


def build_cover_letter(
    text: str,
    *,
    name: str = "",
    email: str = "",
    phone: str = "",
    location: str = "",
    output_path: str | Path,
) -> Path:
    """Build a cover letter DOCX from plain text.

    Returns the output Path.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Header: Name
    if name:
        heading = doc.add_paragraph()
        run = heading.add_run(name)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Calibri"

    # Contact line
    contact_parts = [p for p in [email, phone, location] if p]
    if contact_parts:
        contact = doc.add_paragraph()
        run = contact.add_run(" | ".join(contact_parts))
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    # Date
    date_para = doc.add_paragraph()
    run = date_para.add_run(date.today().strftime("%B %d, %Y"))
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    doc.add_paragraph()  # blank line

    # Body paragraphs
    for paragraph in text.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        p = doc.add_paragraph()
        run = p.add_run(paragraph)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    log.info("Cover letter saved to %s", out)
    return out
