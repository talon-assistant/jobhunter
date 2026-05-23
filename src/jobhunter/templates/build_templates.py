#!/usr/bin/env python3
"""Generate the built-in ATS-friendly resume templates.

Run this script to regenerate the template DOCX files:
    python -m jobhunter.templates.build_templates

Each template is a DOCX with placeholder content that docx_builder.py
replaces with actual resume bullets. The placeholders use a consistent
marker format that the builder recognizes.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

TEMPLATE_DIR = Path(__file__).parent


def _add_horizontal_rule(doc, color: str = "CCCCCC"):
    """Add a thin horizontal rule paragraph.

    *color* is a 6-char hex string like ``"1A1A2E"``.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): color,
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _set_margins(doc, top=0.5, bottom=0.5, left=0.75, right=0.75):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def _add_placeholder_bullet(doc, text: str, font_name: str, font_size: Pt):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    return p


# ──────────────────────────────────────────────────────────────
# Template 1: Classic
# Clean, conservative design. Calibri. Ideal for corporate roles.
# ──────────────────────────────────────────────────────────────

def build_classic():
    doc = Document()
    _set_margins(doc)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{NAME}}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{EMAIL}} | {{PHONE}} | {{LOCATION}}")
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    _add_horizontal_rule(doc, "1A1A2E")

    # Section placeholder pattern
    for section in ["PROFESSIONAL SUMMARY", "EXPERIENCE", "SKILLS", "EDUCATION", "CERTIFICATIONS"]:
        p = doc.add_paragraph()
        run = p.add_run(section)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)

        _add_horizontal_rule(doc, "CCCCCC")

        # Role header (for experience sections)
        if section == "EXPERIENCE":
            p = doc.add_paragraph()
            run = p.add_run("{{ROLE_TITLE}} — {{COMPANY}}")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            p2 = doc.add_paragraph()
            run2 = p2.add_run("{{DATES}} | {{LOCATION}}")
            run2.font.size = Pt(10)
            run2.font.name = "Calibri"
            run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        _add_placeholder_bullet(doc, "{{BULLET}}", "Calibri", Pt(10.5))
        _add_placeholder_bullet(doc, "{{BULLET}}", "Calibri", Pt(10.5))

    doc.save(str(TEMPLATE_DIR / "classic.docx"))
    print("  Built: classic.docx")


# ──────────────────────────────────────────────────────────────
# Template 2: Modern
# Slightly bolder look. Calibri with accent color headers.
# ──────────────────────────────────────────────────────────────

def build_modern():
    doc = Document()
    _set_margins(doc, top=0.6, bottom=0.6, left=0.7, right=0.7)

    ACCENT_RGB = RGBColor(0x2B, 0x57, 0x9A)
    ACCENT_HEX = "2B579A"

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("{{NAME}}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Calibri"
    run.font.color.rgb = ACCENT_RGB

    # Contact on same line area
    p = doc.add_paragraph()
    run = p.add_run("{{EMAIL}} | {{PHONE}} | {{LOCATION}}")
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    _add_horizontal_rule(doc, ACCENT_HEX)

    for section in ["SUMMARY", "EXPERIENCE", "TECHNICAL SKILLS", "EDUCATION", "CERTIFICATIONS"]:
        p = doc.add_paragraph()
        run = p.add_run(section)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        run.font.color.rgb = ACCENT_RGB
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

        _add_horizontal_rule(doc, ACCENT_HEX)

        if section == "EXPERIENCE":
            p = doc.add_paragraph()
            run = p.add_run("{{ROLE_TITLE}}")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            run2 = p.add_run(" — {{COMPANY}}")
            run2.font.size = Pt(11)
            run2.font.name = "Calibri"
            p2 = doc.add_paragraph()
            run3 = p2.add_run("{{DATES}} | {{LOCATION}}")
            run3.font.size = Pt(10)
            run3.font.name = "Calibri"
            run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        _add_placeholder_bullet(doc, "{{BULLET}}", "Calibri", Pt(10.5))
        _add_placeholder_bullet(doc, "{{BULLET}}", "Calibri", Pt(10.5))

    doc.save(str(TEMPLATE_DIR / "modern.docx"))
    print("  Built: modern.docx")


# ──────────────────────────────────────────────────────────────
# Template 3: Executive
# Garamond, understated elegance. For senior/director+ roles.
# ──────────────────────────────────────────────────────────────

def build_executive():
    doc = Document()
    _set_margins(doc, top=0.75, bottom=0.75, left=1.0, right=1.0)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{NAME}}")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Garamond"
    run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
    p.paragraph_format.space_after = Pt(2)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{EMAIL}} | {{PHONE}} | {{LOCATION}}")
    run.font.size = Pt(10)
    run.font.name = "Garamond"
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _add_horizontal_rule(doc, "2C2C2C")

    for section in ["EXECUTIVE SUMMARY", "PROFESSIONAL EXPERIENCE", "CORE COMPETENCIES", "EDUCATION & CERTIFICATIONS"]:
        p = doc.add_paragraph()
        run = p.add_run(section)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Garamond"
        run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)

        _add_horizontal_rule(doc, "999999")

        if section == "PROFESSIONAL EXPERIENCE":
            p = doc.add_paragraph()
            run = p.add_run("{{ROLE_TITLE}}, {{COMPANY}}")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Garamond"
            p2 = doc.add_paragraph()
            run2 = p2.add_run("{{DATES}} | {{LOCATION}}")
            run2.font.size = Pt(10)
            run2.font.name = "Garamond"
            run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run2.italic = True

        _add_placeholder_bullet(doc, "{{BULLET}}", "Garamond", Pt(11))
        _add_placeholder_bullet(doc, "{{BULLET}}", "Garamond", Pt(11))

    doc.save(str(TEMPLATE_DIR / "executive.docx"))
    print("  Built: executive.docx")


# ──────────────────────────────────────────────────────────────
# Template 4: Compact
# Arial, tight spacing. Fits more content on 1-2 pages.
# Good for technical roles with lots of bullet points.
# ──────────────────────────────────────────────────────────────

def build_compact():
    doc = Document()
    _set_margins(doc, top=0.4, bottom=0.4, left=0.6, right=0.6)

    # Name
    p = doc.add_paragraph()
    run = p.add_run("{{NAME}}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Arial"
    p.paragraph_format.space_after = Pt(0)

    # Contact — compact single line
    p = doc.add_paragraph()
    run = p.add_run("{{EMAIL}} | {{PHONE}} | {{LOCATION}}")
    run.font.size = Pt(9)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.space_after = Pt(4)

    _add_horizontal_rule(doc, "333333")

    for section in ["SUMMARY", "EXPERIENCE", "SKILLS", "PROJECTS", "EDUCATION"]:
        p = doc.add_paragraph()
        run = p.add_run(section.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)

        if section == "EXPERIENCE":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run("{{ROLE_TITLE}} — {{COMPANY}}")
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Arial"
            run2 = p.add_run("  ({{DATES}})")
            run2.font.size = Pt(9)
            run2.font.name = "Arial"
            run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        _add_placeholder_bullet(doc, "{{BULLET}}", "Arial", Pt(9.5))
        _add_placeholder_bullet(doc, "{{BULLET}}", "Arial", Pt(9.5))

    doc.save(str(TEMPLATE_DIR / "compact.docx"))
    print("  Built: compact.docx")


def main():
    print("Building resume templates...")
    build_classic()
    build_modern()
    build_executive()
    build_compact()
    print(f"Done. Templates saved to {TEMPLATE_DIR}")


if __name__ == "__main__":
    main()
